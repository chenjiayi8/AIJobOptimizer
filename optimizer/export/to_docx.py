"""
This module provides functions to create a Word document from a given \
template. The main function is `to_docx` which creates a Word document from \
a given template , and replaces placeholders with the corresponding parameters
"""
from docx import Document
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.shared import Pt
import re


def move_table_after(table, paragraph):
    tbl, p = table._tbl, paragraph._p
    p.addnext(tbl)


def delete_paragraph(paragraph):
    p = paragraph._element
    p.getparent().remove(p)
    paragraph._p = paragraph._element = None


def move_paragraph_after(p_before, p_after):
    p_before._p.addnext(p_after._p)


def create_project_name(doc, proj):
    para = doc.add_paragraph()
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.style.font.size = Pt(11)
    para.add_run('Project: ').bold = True
    para.add_run(proj['title'])
    para.paragraph_format.space_after = Pt(0)
    return para


def create_project_description(doc, proj):
    para = doc.add_paragraph()
    para.add_run(proj['description']).italic = True
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.style.font.size = Pt(11)
    para.paragraph_format.space_after = Pt(0)
    return para


def create_contribution_head(doc):
    para = doc.add_paragraph()
    run_head = para.add_run('Key Contributions:')
    run_head.italic = True
    run_head.bold = True
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.style.font.size = Pt(11)
    para.paragraph_format.space_after = Pt(0)
    return para


def render_contribution(para, contribution):
    left_bold = [m.start() for m in re.finditer('<b>', contribution)]
    right_bold = [m.start() for m in re.finditer('</b>', contribution)]
    start = 0
    if len(left_bold) != len(right_bold):
        raise ValueError('tag open and close not matched')
    if len(left_bold) == 0:
        para.add_run(contribution)
        return

    for i in range(len(left_bold)):
        l = left_bold[i]
        r = right_bold[i]
        if start < l:
            para.add_run(contribution[start:l])
        para.add_run(contribution[l+3:r]).bold = True
        start = r + 4
    if right_bold[-1] < len(contribution):
        para.add_run(contribution[right_bold[-1]+4:])


def create_contribution(doc, contribution):
    para = doc.add_paragraph(style='List Bullet 2')
    render_contribution(para, contribution)
    para.alignment = WD_ALIGN_PARAGRAPH.LEFT
    para.style.font.size = Pt(11)
    return para


def create_contributions(doc, proj):
    paras = []
    for contribution in proj['contributions']:
        paras.append(create_contribution(doc, contribution))
    return paras


def create_project(doc, proj):
    paras = []
    paras.append(create_project_name(doc, proj))
    paras.append(create_project_description(doc, proj))
    paras.append(create_contribution_head(doc))
    paras += create_contributions(doc, proj)
    return paras


def create_projects(doc, exp):
    paras = []
    for proj in exp['projects']:
        paras += create_project(doc, proj)

    return paras


def create_company(doc, exp):
    str_range = f"({exp['date_range']})"
    para = doc.add_paragraph()
    para.add_run(exp['title']).bold = True
    para.add_run(' at ')
    para.add_run(exp['company'])
    padding = ' '*round(3)
    para.add_run(padding)
    para.add_run(str_range)
    para.paragraph_format.space_before = Pt(9)
    para.paragraph_format.space_after = Pt(0)
    return para


def write_experiences(doc, para, experiences):
    para_placeholder = para
    paragraphs = []
    for exp in experiences.values():
        para = create_company(doc, exp)
        para_projects = create_projects(doc, exp)
        paragraphs.append(para)
        paragraphs += para_projects

    for p in paragraphs[::-1]:
        move_paragraph_after(para_placeholder, p)
    delete_paragraph(para_placeholder)


def set_paragraph_font_name(para, name):
    for run in para.runs:
        run.font.name = name


def to_docx(input_file, statement, skills_str, experiences, output_file=None, new_font_name='Times New Roman'):
    doc = Document(input_file)
    for para in doc.paragraphs:
        if para.text == '{statement}':
            para.text = statement.strip()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.style.font.size = Pt(11)

        if para.text == '{competencies}':
            para.text = skills_str.strip()
            para.alignment = WD_ALIGN_PARAGRAPH.LEFT
            para.style.font.size = Pt(11)

        if para.text == '{experiences}':
            write_experiences(doc, para, experiences)

    if output_file is None:
        output_file = input_file.split('.')[0] + '_output.docx'

    if new_font_name is not None:
        for para in doc.paragraphs:
            set_paragraph_font_name(para, new_font_name)

    doc.save(output_file)
    return output_file

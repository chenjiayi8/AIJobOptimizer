"""
This module provides functions to create a Word document from a given \
template. The main function is `to_docx` which creates a Word document from \
a given template , and replaces placeholders with the corresponding parameters
"""
from io import BytesIO
import copy
import re
from docx import Document
from docx.enum.text import WD_PARAGRAPH_ALIGNMENT, WD_LINE_SPACING
from docx.shared import Pt, Inches
from simplify_docx import simplify


def move_table_after(table, paragraph):
    """
    Moves the table after the specified paragraph.

    Args:
        table (docx.table.Table): The table object to move.
        paragraph (docx.text.paragraph.Paragraph): The paragraph object to \
            move the table after.
    """
    tbl, para = table._tbl, paragraph._p
    para.addnext(tbl)


def delete_paragraph(paragraph):
    """
    This function takes a paragraph object as input and removes it from the parent element.

    Parameters:
    paragraph (object): The paragraph object to be deleted.\n    \n    Returns:
    None
    """
    para = paragraph._element
    para.getparent().remove(para)
    paragraph._p = paragraph._element = None


def move_paragraph_after(p_before, p_after):
    """Move the paragraph object p_before after p_after.

    Args:
        p_before: The paragraph object to be moved.
        p_after: The paragraph object after which p_before will be moved.

    Returns: None.
    """
    p_before._p.addnext(p_after._p)


def create_project_name(doc, proj):
    """
    Adds a project title to a Word document.

    Args:
        doc: A Word document object.
        proj: A dictionary containing the project details, including the title.

    Returns:
        A Word paragraph object containing the project title.
    """
    para = doc.add_paragraph()
    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    para.style.font.size = Pt(11)
    para.add_run("Project: ").bold = True
    para.add_run(proj["title"])
    para.paragraph_format.space_after = Pt(0)
    return para


def create_project_description(doc, proj):
    """
    Adds a project description to a Word document.

    Args:
        doc: A Word document object.
        proj: A dictionary containing the project details, including the \
        description.

    Returns:
        A Word paragraph object containing the project description.
    """
    para = doc.add_paragraph()
    para.add_run(proj["description"]).italic = True
    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    para.style.font.size = Pt(11)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.15
    return para


def create_contribution_head(doc):
    """
    Create a paragraph containing the header for Key Contributions section.

    Parameters:
    doc (docx.document.Document): the Document instance to which the header \
    paragraph is to be added.

    Returns:
    docx.text.paragraph.Paragraph: the created Paragraph instance containing the header.
    """
    para = doc.add_paragraph()
    run_head = para.add_run("Key Contributions:")
    run_head.italic = True
    run_head.bold = True
    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    para.style.font.size = Pt(11)
    para.paragraph_format.space_after = Pt(0)
    return para


def render_contribution(para, contribution):
    """
    Renders a contribution onto a paragraph object. A contribution is a \
    string with optional HTML-like `<b>` tags to indicate bold text. The \
    function inserts runs of text into the paragraph, alternating between \
    bold and non-bold text according to the `<b>` tags.

    Args:
        para (docx.text.paragraph.Paragraph): The paragraph object to insert \
        the contribution into.
        contribution (str): The contribution to render. May include `<b>` \
        tags to indicate bold text.

    Returns:
        None

    Raises:
        ValueError: If the number of `<b>` tags does not match the number of \
        `</b>` tags.
    """
    left_bold = [m.start() for m in re.finditer("<b>", contribution)]
    right_bold = [m.start() for m in re.finditer("</b>", contribution)]
    start = 0
    if len(left_bold) != len(right_bold):
        raise ValueError("tag open and close not matched")
    if len(left_bold) == 0:
        para.add_run(contribution)
        return

    for left, right in zip(left_bold, right_bold):
        if start < left:
            para.add_run(contribution[start:left])
        para.add_run(contribution[left + 3 : right]).bold = True
        start = right + 4
    if right_bold[-1] < len(contribution):
        para.add_run(contribution[right_bold[-1] + 4 :])


def add_bullet_point(para):
    """
    This function takes a paragraph object 'para' and formats it with a \
    bullet point character.

    Parameters:
    para: Paragraph - A single paragraph object from the python-docx library

    Returns:
    None
    """
    # Set bullet properties
    para.paragraph_format.left_indent = Inches(0.5)
    para.paragraph_format.first_line_indent = Inches(-0.25)
    para.paragraph_format.tab_stops.add_tab_stop(Inches(0.5))
    # Replace 'Symbol' with the desired font
    para.style.font.name = "Times New Roman"
    para.text = "\u25CF    " + para.text  # Prepend bullet point character


def create_contribution(doc, contribution):
    """
    Add a bulleted contribution to a given document.

    Arguments:
    - doc: An instance of docx.Document class representing the document to add contribution to.
    - contribution: A string representing the text to be added as contribution.

    Returns:
    - An instance of docx.text.paragraph.Paragraph class representing the added paragraph.
    """
    para = doc.add_paragraph()
    add_bullet_point(para)
    render_contribution(para, contribution)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    para.style.font.size = Pt(11)
    para.paragraph_format.space_before = Pt(0)
    para.paragraph_format.space_after = Pt(0)
    para.paragraph_format.line_spacing = 1.15
    return para


def create_contributions(doc, proj):
    """
    Creates a list of contributions for a project and adds them to a document.

    Args:
    doc (Document): The document to which the contributions will be added.
    proj (dict): A dictionary representing a project.

    Returns:
    list: A list of paragraphs representing the contributions added to the \
    document.
    """
    paras = []
    for contribution in proj["contributions"]:
        paras.append(create_contribution(doc, contribution))
    return paras


def create_project(doc, proj):
    """
    Creates and returns a list of paragraphs that describe a project.

    Args:
        doc (document): The document object representing the project.
        proj (project): The project object containing metadata.

    Returns:
        list: A list of paragraphs describing the project.
    """
    paras = []
    paras.append(create_project_name(doc, proj))
    paras.append(create_project_description(doc, proj))
    paras.append(create_contribution_head(doc))
    paras += create_contributions(doc, proj)
    return paras


def create_projects(doc, exp):
    """
    Create and return a list of paragraph elements in a given document (doc) \
    based on information about projects from a given experiment (exp).

    This function loops through each project in exp['projects'] and calls the \
    create_project function for each one, collecting the resulting paragraphs \
    into a list. The final list of paragraphs is returned.

    Args:
        doc: A document object where paragraphs will be created.
        exp: A dictionary containing information about the experiment, \
            including a list of projects under the 'projects' key.

    Returns:
        A list of paragraph elements to be added to the given document.
    """
    paras = []
    for proj in exp["projects"]:
        paras += create_project(doc, proj)

    return paras


def create_company(doc, exp):
    """
    Creates a company paragraph in a Word document.

    Parameters:
    -----------
    doc : Document
        The Word document to add the paragraph to.
    exp : dict
        The experience dictionary containing the following keys:
            - title : str
                The job title.
            - company : str
                The name of the company.
            - date_range : str
                The date range of the job.

    Returns:
    --------
    para : Paragraph
        The created paragraph.
    """
    str_range = f"({exp['date_range']})"
    para = doc.add_paragraph()
    para.add_run(exp["title"]).bold = True
    para.add_run(" at ")
    para.add_run(exp["company"])
    padding = " " * round(3)
    para.add_run(padding)
    para.add_run(str_range)
    para.paragraph_format.space_before = Pt(9)
    para.paragraph_format.space_after = Pt(0)
    return para


def write_experiences(doc, para, experiences):
    """
    Insert a list of experiences under a given paragraph in a Word document.

    Parameters:
        doc (docx.Document): The document to insert the experiences into.
        para (docx.Paragraph): The paragraph to insert the experiences under.
        experiences (dict): A dictionary containing experience data.

    Returns:
        None
    """
    para_placeholder = para
    paragraphs = []
    for exp in experiences:
        para = create_company(doc, exp)
        para_projects = create_projects(doc, exp)
        paragraphs.append(para)
        paragraphs += para_projects

    for paragraph in paragraphs[::-1]:
        move_paragraph_after(para_placeholder, paragraph)
    delete_paragraph(para_placeholder)


def set_paragraph_font_name(para, name):
    """
    Sets the font name for all runs within a given paragraph.

    Parameters:
    para (docx.text.paragraph.Paragraph): The paragraph whose font name is to \
    be set.
    name (str): The name of the font to be set.

    Returns:
    None
    """
    for run in para.runs:
        run.font.name = name


def create_docx(bytes_data):
    """
    Create a new Word document from the given bytes.

    Args:
        bytes_data (bytes): The raw bytes representing the Word document.

    Returns:
        docx.Document: A new `docx.Document` object constructed from the given bytes.

    """
    return Document(BytesIO(copy.deepcopy(bytes_data)))


def validate_template(bytes_data, template_fields):
    """
    Given a byte stream of a docx file as bytes_data and a list of
    strings representing template fields as template_fields, this
    function checks if all template fields are present in the docx
    file. If a template field is found, it is removed from the list
    of targets until all targets have been found or not.

    Args:
    - bytes_data: a byte stream of a docx file
    - template_fields: a list of strings representing template fields

    Returns:
    - a list of strings representing the unfound template fields,
      which will be an empty list if all targets have been found.
    """
    targets = copy.deepcopy(template_fields)
    doc = create_docx(bytes_data)
    for para in doc.paragraphs:
        if para.text in targets:
            targets.remove(para.text)

    return targets


def write_letter(letter, new_font_name="Times New Roman"):
    """
    Creates a Word document consisting of a single paragraph with the given \
    letter content. The font size of the paragraph is set to 12pt, alignment \
    is left and line spacing is single. The font name of the paragraph is set \
    to the provided font name. If no font name is provided, the default value \
    is 'Times New Roman'.

    Args:
        letter (str): The content of the letter to be added to the Word \
        document.
        new_font_name (str, optional): The font name to be used for the \
        paragraph. Defaults to 'Times New Roman'.

    Returns:
        bytes: The contents of the saved Word document as bytes.
    """
    doc = Document()
    para = doc.add_paragraph(letter)
    para.style.font.size = Pt(12)
    para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
    para.paragraph_format.line_spacing = 1.15
    for para in doc.paragraphs:
        set_paragraph_font_name(para, new_font_name)
    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.read()


def to_docx(
    bytes_data,
    statement,
    skills_str,
    experiences,
    new_font_name="Times New Roman",
):
    """
    Create a .docx file from bytes data.

    Parameters:
        bytes_data (bytes): The bytes data to convert to a .docx file.
        statement (str): The statement to insert into the .docx file.
        skills_str (str): The competencies to insert into the .docx file.
        experiences (Optional[List[Dict[str,str]]]): A list of dictionaries \
        representing experiences to insert into the .docx file.
        new_font_name (Optional[str]): The name of the font to use for the \
        .docx file. Defaults to 'Times New Roman'.

    Returns:
        bytes: The bytes representation of the .docx file.
    """
    doc = create_docx(bytes_data)
    for para in doc.paragraphs:
        if para.text == "{statement}":
            para.text = statement.strip()
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.style.font.size = Pt(11)
            para.paragraph_format.line_spacing = 1.15

        if para.text == "{competencies}":
            para.text = skills_str.strip()
            para.alignment = WD_PARAGRAPH_ALIGNMENT.LEFT
            para.style.font.size = Pt(11)
            para.paragraph_format.line_spacing = 1.15

        if para.text == "{experiences}":
            write_experiences(doc, para, experiences)

    if new_font_name is not None:
        for para in doc.paragraphs:
            set_paragraph_font_name(para, new_font_name)

    file_stream = BytesIO()
    doc.save(file_stream)
    file_stream.seek(0)
    return file_stream.read()


def extract_text_from_docx(doc):
    """
    Extracts plain text from a docx file.

    Args:
        doc: The docx file to extract plain text from.

    Returns:
        A string containing the plain text extracted from the docx file.
    """

    def extract_text(json_obj):
        text = []

        if isinstance(json_obj, dict):
            obj_type = json_obj.get("TYPE", "")
            obj_value = json_obj.get("VALUE", [])

            if obj_type == "text":
                text.append(obj_value)
            elif isinstance(obj_value, (list, dict)):
                text.extend(extract_text(obj_value))

        elif isinstance(json_obj, list):
            for item in json_obj:
                text.extend(extract_text(item))

        return text

    json_obj = simplify(doc)
    result = extract_text(json_obj)
    text = "\n".join(result)
    return text


def docx_to_text(bytes_data):
    """
    Takes a byte string as input, which represents the binary content of a \
    docx file.
    Returns a string that contains the plain text content of the document.

    Parameters:
    bytes_data (bytes): A byte string that represents the binary content of a \
    docx file.

    Returns:
    str: A string that contains the plain text content of the input docx file.
    """
    doc = create_docx(bytes_data)
    result = extract_text_from_docx(doc)
    return result
